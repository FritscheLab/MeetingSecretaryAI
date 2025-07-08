import json
import os
from datetime import datetime, date
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
import argparse

today_str = date.today().strftime("%Y-%m-%d")

def add_bold_text(paragraph, bold_text, normal_text):
    run = paragraph.add_run(bold_text)
    run.bold = True
    paragraph.add_run(normal_text)

def format_month_year(date_str):
    """
    Converts 'YYYY-MM-DD' into 'Month YYYY'. If invalid, returns the original string.
    """
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        return date_obj.strftime('%B %Y')
    except ValueError:
        return date_str

def format_date(date_str):
    """
    Converts 'YYYY-MM-DD' into 'Month DD, YYYY'. Returns "TBD" if applicable.
    """
    if date_str.upper() == "TBD":
        return "TBD"
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        return date_obj.strftime("%B %d, %Y")
    except ValueError:
        return date_str

def format_time(time_str):
    """
    Converts 'HH:MM' into 'HH:MM AM/PM'. Returns the original string if invalid.
    """
    try:
        time_obj = datetime.strptime(time_str, "%H:%M")
        return time_obj.strftime("%I:%M %p")
    except ValueError:
        return time_str

def set_background_color(run, color):
    """
    Highlights text background with a given hex color (e.g. 'FFFF00' for yellow).
    """
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    rPr.append(shd)

def add_bookmark(paragraph, bookmark_name):
    """
    Inserts a bookmark start and end into a paragraph.
    Used for internal links within the document.
    """
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), bookmark_name)
    paragraph._p.insert(0, bookmark_start)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')
    paragraph._p.append(bookmark_end)

def add_internal_hyperlink(paragraph, bookmark_name, text, color="0000FF", underline=True):
    """
    Creates an internal hyperlink to a bookmark within the Word doc.
    """
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def generate_action_items_recap(json_data):
    """
    Generate a recap of all action items consolidated by unique responsible party.
    It traverses sections and subsections and aggregates tasks for each responsible.
    """
    recap = {}
    for section in json_data.get('sections', []):
        # Process action_items at the section level
        for item in section.get('action_items', []):
            if isinstance(item, dict):
                responsible = item.get('responsible', '').strip()
                task = item.get('task', '').strip()
                if responsible and task:
                    recap.setdefault(responsible, []).append(task)
        # Process action_items within subsections
        for subsection in section.get('subsections', []):
            for item in subsection.get('action_items', []):
                if isinstance(item, dict):
                    responsible = item.get('responsible', '').strip()
                    task = item.get('task', '').strip()
                    if responsible and task:
                        recap.setdefault(responsible, []).append(task)
    return recap

def create_meeting_minutes(json_data, include_rationale=False, include_recommendations=False):
    doc = Document()

    # Set base styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(3)

    title_style = doc.styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(14)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_style.paragraph_format.space_before = Pt(6)
    title_style.paragraph_format.space_after = Pt(6)

    # Adjust heading styles
    for i in range(1, 7):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(11)
        heading_style.paragraph_format.space_before = Pt(3)
        heading_style.paragraph_format.space_after = Pt(3)

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title
    title_text = json_data['meeting_details']['title']
    title = doc.add_paragraph(title_text)
    title.style = 'Title'

    # Meeting Details: Date, Time, Location
    meeting_details = json_data['meeting_details']
    details_paragraph = doc.add_paragraph(style='Normal')
    add_bold_text(details_paragraph, "Date: ", format_date(meeting_details['date']) + "\n")
    add_bold_text(details_paragraph, "Time: ", format_time(meeting_details['time']) + "\n")
    add_bold_text(details_paragraph, "Location: ", meeting_details['location'] + "\n")

    # Attendees List
    def valid_person(person):
        return bool(person.get('first_name')) and bool(person.get('last_name'))
    
    attendee_list = []
    for person in meeting_details['attendees']:
        if valid_person(person):
            attendee_list.append(f"{person['first_name']} {person['last_name']}")
    if attendee_list:
        attendees_paragraph = doc.add_paragraph(style='Normal')
        add_bold_text(attendees_paragraph, "Attendees: ", ", ".join(attendee_list))
    
    # Absentees List
    absentee_list = []
    for person in meeting_details['absentees']:
        if valid_person(person):
            absentee_list.append(f"{person['first_name']} {person['last_name']}")
    if absentee_list:
        absentees_paragraph = doc.add_paragraph(style='Normal')
        add_bold_text(absentees_paragraph, "Absentees: ", ", ".join(absentee_list))
    
    # Link to Summarized Action Items at the very front
    link_paragraph = doc.add_paragraph()
    link_paragraph.paragraph_format.space_before = Pt(12)
    link_paragraph.paragraph_format.space_after = Pt(12)
    add_internal_hyperlink(link_paragraph, "SummarizedActionItems", "Jump to Summarized Action Items")
    
    # Reflection Section
    doc.add_heading('Reflection', level=1)
    doc.add_paragraph(json_data.get('reflection', ''), style='Normal')

    # Process each Section
    for section in json_data['sections']:
        doc.add_heading(section['title'], level=1)

        # Discussion Points for the section
        if section.get('discussion_points'):
            doc.add_paragraph("Discussion Points:", style='List Bullet')
            for point in section['discussion_points']:
                doc.add_paragraph(point, style='List Bullet 2')

        # Decisions for the section
        if section.get('decisions'):
            doc.add_paragraph("Decisions:", style='List Bullet')
            for decision in section['decisions']:
                doc.add_paragraph(decision, style='List Bullet 2')

        # Action Items for the section (with responsible person highlighted)
        if section.get('action_items'):
            doc.add_paragraph("Action Items:", style='List Bullet')
            for item in section['action_items']:
                paragraph = doc.add_paragraph(style='List Bullet 2')
                if isinstance(item, dict):
                    task_text = item.get('task', '')
                    if task_text:
                        paragraph.add_run(task_text)
                    responsible = item.get('responsible')
                    if responsible:
                        paragraph.add_run(" (")
                        run_responsible = paragraph.add_run(f"{responsible}")
                        set_background_color(run_responsible, "FFFF00")  # highlight responsible party
                        paragraph.add_run(")")
                else:
                    paragraph.add_run(str(item))

        # Append Reference Titles as inline hyperlinks (if any)
        if section.get('reference_titles'):
            ref_paragraph = doc.add_paragraph(style='List Bullet')
            ref_paragraph.add_run("Ref: ")
            for i, ref_title in enumerate(section['reference_titles']):
                if i > 0:
                    ref_paragraph.add_run(", ")
                add_internal_hyperlink(ref_paragraph, "ReferenceAppendix", ref_title)

        # Key Recommendations for the section (conditional)
        if include_recommendations and section.get('key_recommendations'):
            doc.add_paragraph("Key Recommendations:", style='List Bullet')
            for rec in section['key_recommendations']:
                doc.add_paragraph(rec, style='List Bullet 2')

        # Rationale for the section (conditional)
        if include_rationale and section.get('rationale'):
            doc.add_paragraph("Rationale:", style='List Bullet')
            for rationale in section['rationale']:
                doc.add_paragraph(rationale, style='List Bullet 2')

        # Process Subsections (if any)
        if section.get('subsections'):
            for subsection in section['subsections']:
                doc.add_heading(subsection['title'], level=2)
                if subsection.get('discussion_points'):
                    doc.add_paragraph("Discussion Points:", style='List Bullet')
                    for point in subsection['discussion_points']:
                        doc.add_paragraph(point, style='List Bullet 2')
                if subsection.get('decisions'):
                    doc.add_paragraph("Decisions:", style='List Bullet')
                    for decision in subsection['decisions']:
                        doc.add_paragraph(decision, style='List Bullet 2')
                if subsection.get('action_items'):
                    doc.add_paragraph("Action Items:", style='List Bullet')
                    for item in subsection['action_items']:
                        paragraph = doc.add_paragraph(style='List Bullet 2')
                        if isinstance(item, dict):
                            task_text = item.get('task', '')
                            if task_text:
                                paragraph.add_run(task_text)
                            responsible = item.get('responsible')
                            if responsible:
                                paragraph.add_run(" (")
                                run_responsible = paragraph.add_run(f"{responsible}")
                                set_background_color(run_responsible, "FFFF00")
                                paragraph.add_run(")")
                        else:
                            paragraph.add_run(str(item))
                if include_recommendations and subsection.get('key_recommendations'):
                    doc.add_paragraph("Key Recommendations:", style='List Bullet')
                    for rec in subsection['key_recommendations']:
                        doc.add_paragraph(rec, style='List Bullet 2')
                if include_rationale and subsection.get('rationale'):
                    doc.add_paragraph("Rationale:", style='List Bullet')
                    for rationale in subsection['rationale']:
                        doc.add_paragraph(rationale, style='List Bullet 2')
                # Append Reference Titles for subsection inline
                if subsection.get('reference_titles'):
                    ref_paragraph = doc.add_paragraph(style='List Bullet')
                    ref_paragraph.add_run("References: ")
                    for i, ref_title in enumerate(subsection['reference_titles']):
                        if i > 0:
                            ref_paragraph.add_run(", ")
                        add_internal_hyperlink(ref_paragraph, "ReferenceAppendix", ref_title)

    # Next Meeting and Adjournment Section
    doc.add_heading('Next Meeting and Adjournment', level=1)
    next_meeting = json_data.get('next_meeting', {})
    adjournment = json_data.get('adjournment', {})

    date_val = next_meeting.get('date', '').strip()
    time_val = next_meeting.get('time', '').strip()
    location_val = next_meeting.get('location', '').strip()

    if date_val and time_val and location_val:
        next_meeting_text = f"Next Meeting: {format_date(date_val)} at {format_time(time_val)}, {location_val}."
    else:
        next_meeting_text = "Next Meeting: Information not available."

    adjournment_time = adjournment.get('time', '').strip()
    if adjournment_time:
        adjournment_text = f"Adjournment: {format_time(adjournment_time)}."
    else:
        adjournment_text = "Adjournment: Not specified."

    doc.add_paragraph(f"{next_meeting_text}\n{adjournment_text}", style='Normal')

    # Summarized Action Items Section (inserted on a new page)
    doc.add_page_break()
    target_paragraph = doc.add_heading('Summarized Action Items', level=1)
    add_bookmark(target_paragraph, "SummarizedActionItems")

    recap_dict = generate_action_items_recap(json_data)

    # create table
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Responsible'
    hdr_cells[1].text = 'Tasks'
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[1].paragraphs[0].runs[0].bold = True

    # populate rows using recap_dict
    for responsible, tasks in recap_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = responsible

        task_cell = row_cells[1]
        # clear default content
        task_cell.paragraphs[0].text = ''
        # add each task as a bullet
        for task in tasks:
            p = task_cell.add_paragraph(style='List Bullet')
            p.add_run(task)

    # Reference Appendix Section (inserted on a new page)
    doc.add_page_break()
    ref_heading = doc.add_heading("Reference Appendix", level=1)
    add_bookmark(ref_heading, "ReferenceAppendix")
    if json_data.get('references'):
        for ref in json_data['references']:
            doc.add_paragraph(f"Title: {ref.get('title', '')}", style='List Bullet')
            doc.add_paragraph(f"Type: {ref.get('reference_type', '')}", style='List Bullet 2')
            doc.add_paragraph(f"Identifier: {ref.get('identifier', 'N/A')}", style='List Bullet 2')
            doc.add_paragraph(f"Description: {ref.get('description', '')}", style='List Bullet 2')
    else:
        doc.add_paragraph("No references available.", style='Normal')

    # Footer with Version Info and Page Numbers
    footer = doc.sections[0].footer
    version_paragraph = footer.add_paragraph()
    version_paragraph.text = f"Version: 1.0 ({format_date(today_str)})"
    version_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    page_paragraph = footer.add_paragraph()
    page_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = page_paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)

    return doc

def create_meeting_minutes_markdown(json_data, include_rationale=False, include_recommendations=False):
    """Generate Markdown formatted meeting minutes from JSON data."""
    lines = []
    # Document Title
    title_text = json_data['meeting_details']['title']
    lines.append(f"# {title_text}")
    lines.append("")

    # Meeting Details
    meeting_details = json_data['meeting_details']
    lines.append(f"**Date:** {format_date(meeting_details['date'])}")
    lines.append(f"**Time:** {format_time(meeting_details['time'])}")
    lines.append(f"**Location:** {meeting_details['location']}")
    lines.append("")

    def valid_person(person):
        return bool(person.get('first_name')) and bool(person.get('last_name'))
    
    attendee_list = []
    for person in meeting_details['attendees']:
        if valid_person(person):
            attendee_list.append(f"{person['first_name']} {person['last_name']}")
    if attendee_list:
        lines.append(f"**Attendees:** {', '.join(attendee_list)}")
    
    absentee_list = []
    for person in meeting_details['absentees']:
        if valid_person(person):
            absentee_list.append(f"{person['first_name']} {person['last_name']}")
    if absentee_list:
        lines.append(f"**Absentees:** {', '.join(absentee_list)}")
    lines.append("")

    # Reflection Section
    lines.append("## Reflection")
    lines.append(json_data.get('reflection', ''))
    lines.append("")

    # Process each Section
    for section in json_data['sections']:
        lines.append(f"## {section['title']}")
        lines.append("")
        if section.get('discussion_points'):
            lines.append("**Discussion Points:**")
            for point in section['discussion_points']:
                lines.append(f"- {point}")
            lines.append("")
        if section.get('decisions'):
            lines.append("**Decisions:**")
            for decision in section['decisions']:
                lines.append(f"- {decision}")
            lines.append("")
        if section.get('action_items'):
            lines.append("**Action Items:**")
            for item in section['action_items']:
                if isinstance(item, dict):
                    task_text = item.get('task', '')
                    responsible = item.get('responsible', '')
                    if responsible:
                        lines.append(f"- {task_text} (Responsible: **{responsible}**)")
                    else:
                        lines.append(f"- {task_text}")
                else:
                    lines.append(f"- {item}")
            lines.append("")
        if include_recommendations and section.get('key_recommendations'):
            lines.append("**Key Recommendations:**")
            for rec in section['key_recommendations']:
                lines.append(f"- {rec}")
            lines.append("")
        if include_rationale and section.get('rationale'):
            lines.append("**Rationale:**")
            for rationale in section['rationale']:
                lines.append(f"- {rationale}")
            lines.append("")
        # Append reference titles inline as hyperlinks (if any)
        if section.get('reference_titles'):
            refs = ", ".join(f"[{title}](#reference-appendix)" for title in section['reference_titles'])
            lines.append(f"Ref: {refs}")
            lines.append("")
        # Process Subsections
        if section.get('subsections'):
            for subsection in section['subsections']:
                lines.append(f"### {subsection['title']}")
                lines.append("")
                if subsection.get('discussion_points'):
                    lines.append("**Discussion Points:**")
                    for point in subsection['discussion_points']:
                        lines.append(f"- {point}")
                    lines.append("")
                if subsection.get('decisions'):
                    lines.append("**Decisions:**")
                    for decision in subsection['decisions']:
                        lines.append(f"- {decision}")
                    lines.append("")
                if subsection.get('action_items'):
                    lines.append("**Action Items:**")
                    for item in subsection['action_items']:
                        if isinstance(item, dict):
                            task_text = item.get('task', '')
                            responsible = item.get('responsible', '')
                            if responsible:
                                lines.append(f"- {task_text} (Responsible: **{responsible}**)")
                            else:
                                lines.append(f"- {task_text}")
                        else:
                            lines.append(f"- {item}")
                    lines.append("")
                if include_recommendations and subsection.get('key_recommendations'):
                    lines.append("**Key Recommendations:**")
                    for rec in subsection['key_recommendations']:
                        lines.append(f"- {rec}")
                    lines.append("")
                if include_rationale and subsection.get('rationale'):
                    lines.append("**Rationale:**")
                    for rationale in subsection['rationale']:
                        lines.append(f"- {rationale}")
                    lines.append("")
                if subsection.get('reference_titles'):
                    refs = ", ".join(f"[{title}](#reference-appendix)" for title in subsection['reference_titles'])
                    lines.append(f"References: {refs}")
                    lines.append("")
    # Next Meeting and Adjournment Section
    lines.append("## Next Meeting and Adjournment")
    next_meeting = json_data.get('next_meeting', {})
    adjournment = json_data.get('adjournment', {})
    
    date_val = next_meeting.get('date', '').strip()
    time_val = next_meeting.get('time', '').strip()
    location_val = next_meeting.get('location', '').strip()
    
    if date_val and time_val and location_val:
        lines.append(f"**Next Meeting:** {format_date(date_val)} at {format_time(time_val)}, {location_val}.")
    else:
        lines.append("**Next Meeting:** Information not available.")
    
    adjournment_time = adjournment.get('time', '').strip()
    if adjournment_time:
        lines.append(f"**Adjournment:** {format_time(adjournment_time)}.")
    else:
        lines.append("**Adjournment:** Not specified.")
    lines.append("")
    
    # Generate Summarized Action Items Recap from the JSON input    
    lines.append("## Summarized Action Items")
    lines.append("")
    lines.append("| Responsible | Tasks |")
    lines.append("| ----------- | ----- |")

    # use recap_dict directly
    recap_dict = generate_action_items_recap(json_data)
    for responsible, tasks in recap_dict.items():
        # render each task as a markdown bullet list inside the cell using <br> line breaks
        bullets = "<br>".join(f"- {t}" for t in tasks)
        lines.append(f"| {responsible} | {bullets} |")

    lines.append("")
        
    # Reference Appendix Section with Markdown anchor
    lines.append("## Reference Appendix {#reference-appendix}")
    lines.append("")
    if json_data.get('references'):
        for ref in json_data['references']:
            lines.append(f"- **Title:** {ref.get('title', '')}")
            lines.append(f"  - **Type:** {ref.get('reference_type', '')}")
            lines.append(f"  - **Identifier:** {ref.get('identifier', 'N/A')}")
            lines.append(f"  - **Description:** {ref.get('description', '')}")
    else:
        lines.append("No references available.")
    lines.append("")
    
    return "\n".join(lines)

def main(input_json, output_dir, output_prefix, output_format, include_rationale, include_recommendations):
    with open(input_json, 'r') as json_file:
        json_data = json.load(json_file)
    
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate DOCX if needed
    if output_format in ("docx", "both"):
        docx_path = os.path.join(output_dir, f"{output_prefix}.docx")
        document = create_meeting_minutes(json_data, include_rationale=include_rationale, include_recommendations=include_recommendations)
        document.save(docx_path)
        print(f"DOCX document saved to {docx_path}")
    
    # Generate Markdown if needed
    if output_format in ("md", "both"):
        md_path = os.path.join(output_dir, f"{output_prefix}.md")
        markdown_text = create_meeting_minutes_markdown(json_data, include_rationale=include_rationale, include_recommendations=include_recommendations)
        with open(md_path, 'w') as md_file:
            md_file.write(markdown_text)
        print(f"Markdown document saved to {md_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Convert JSON meeting minutes to formatted DOCX and/or Markdown documents."
    )
    parser.add_argument("--input_json", required=True,
                        help="Path to the input JSON file containing meeting minutes.")
    parser.add_argument("--output_dir", default=".",
                        help="Directory to save the output files (default: current directory).")
    parser.add_argument("--output_prefix", default="meeting_minutes",
                        help="Prefix for the output file names (default: meeting_minutes).")
    parser.add_argument("--output_format", choices=["docx", "md", "both"], default="both",
                        help="Output format: 'docx', 'md', or 'both' (default: both).")
    parser.add_argument("--include_rationale", action="store_true",
                        help="Include the rationale sections in the minutes.")
    parser.add_argument("--include_recommendations", action="store_true",
                        help="Include the key recommendations sections in the minutes.")
    args = parser.parse_args()
    main(args.input_json, args.output_dir, args.output_prefix, args.output_format,
         include_rationale=args.include_rationale, include_recommendations=args.include_recommendations)
